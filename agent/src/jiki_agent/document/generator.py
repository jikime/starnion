"""Document generation: create documents in various formats."""

from io import BytesIO
from pathlib import Path

# Bundled Noto Sans KR font for Korean/CJK PDF support.
_FONT_DIR = Path(__file__).parent / "fonts"
_NOTO_SANS_KR = _FONT_DIR / "NotoSansKR.ttf"

# ReportLab font registration (module-level, runs once).
_FONT_REGISTERED = False


def _ensure_font() -> str:
    """Register Noto Sans KR with ReportLab and return the font name."""
    global _FONT_REGISTERED  # noqa: PLW0603
    if not _FONT_REGISTERED and _NOTO_SANS_KR.exists():
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        pdfmetrics.registerFont(TTFont("NotoSansKR", str(_NOTO_SANS_KR)))
        _FONT_REGISTERED = True
    return "NotoSansKR" if _FONT_REGISTERED else "Helvetica"


def generate_pdf(title: str, content: str) -> bytes:
    """Generate a PDF document with full Korean/CJK support using ReportLab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    font_name = _ensure_font()
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        "KorTitle", fontName=font_name, fontSize=18, leading=24, spaceAfter=12,
    ))
    styles.add(ParagraphStyle(
        "KorBody", fontName=font_name, fontSize=11, leading=16, spaceAfter=8,
    ))

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=20 * mm, bottomMargin=20 * mm,
        leftMargin=20 * mm, rightMargin=20 * mm,
    )

    story: list = [Paragraph(title, styles["KorTitle"]), Spacer(1, 6 * mm)]
    for para in content.split("\n\n"):
        text = para.strip()
        if text:
            # Replace single newlines with <br/> for ReportLab.
            story.append(Paragraph(text.replace("\n", "<br/>"), styles["KorBody"]))

    doc.build(story)
    return buf.getvalue()


def generate_docx(title: str, content: str) -> bytes:
    """Generate a DOCX document from text content."""
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    for paragraph in content.split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_xlsx(headers: list[str], rows: list[list]) -> bytes:
    """Generate an XLSX spreadsheet from tabular data."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(headers)
    for row in rows:
        ws.append(row)
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def generate_md(title: str, content: str) -> bytes:
    """Generate a Markdown document."""
    text = f"# {title}\n\n{content}"
    return text.encode("utf-8")


def generate_txt(content: str) -> bytes:
    """Generate a plain text document."""
    return content.encode("utf-8")
